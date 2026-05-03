"""File parsers: take raw bytes, return (kind, content_text, metadata).

content_text is a plain-text rendering used as input to the summarizer and
as a default fallback for tools. It is capped — for huge files the agent
should use the python_sandbox tool to load the original.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document as _Docx
from pypdf import PdfReader


# Hard cap on the plaintext we keep in memory / prompt context.
MAX_TEXT_CHARS = 60_000


_EXT_TO_KIND = {
    ".csv":  "csv",
    ".tsv":  "csv",
    ".xlsx": "xlsx",
    ".xls":  "xlsx",
    ".pdf":  "pdf",
    ".docx": "docx",
    ".doc":  "docx",
    ".txt":  "txt",
    ".md":   "txt",
    ".json": "txt",
    ".log":  "txt",
}


def kind_for(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    return _EXT_TO_KIND.get(ext, "txt")


def _truncate(text: str, limit: int = MAX_TEXT_CHARS) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + f"\n\n[... truncated, {len(text) - limit} more chars not shown ...]"


def parse_csv(raw: bytes, filename: str) -> tuple[str, dict]:
    df = pd.read_csv(io.BytesIO(raw))
    return _df_to_text(df, filename, sheet_name=None)


def parse_xlsx(raw: bytes, filename: str) -> tuple[str, dict]:
    """Multi-sheet aware. Renders each sheet's head + describe, concatenated."""
    xls = pd.ExcelFile(io.BytesIO(raw))
    parts: list[str] = []
    sheets_meta: list[dict] = []
    for name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=name)
        text, meta = _df_to_text(df, filename, sheet_name=name)
        parts.append(f"## Sheet: {name}\n{text}")
        sheets_meta.append({"name": name, **meta})
    return _truncate("\n\n".join(parts)), {"sheets": sheets_meta}


def _df_to_text(df: pd.DataFrame, filename: str, sheet_name: Optional[str]) -> tuple[str, dict]:
    rows, cols = df.shape
    # Render head + describe for the LLM. Keep both bounded.
    head = df.head(20).to_string(max_cols=20, max_colwidth=60)
    try:
        describe = df.describe(include="all").to_string(max_cols=10, max_colwidth=40)
    except Exception:
        describe = "(describe() unavailable)"
    body = (
        f"# {filename}" + (f" — sheet '{sheet_name}'" if sheet_name else "") + "\n"
        f"Shape: {rows} rows × {cols} columns\n"
        f"Columns: {list(df.columns)}\n\n"
        f"### head(20)\n{head}\n\n"
        f"### describe()\n{describe}"
    )
    meta = {
        "rows": int(rows),
        "cols": int(cols),
        "columns": [str(c) for c in df.columns],
        "dtypes": {str(c): str(t) for c, t in df.dtypes.items()},
    }
    return _truncate(body), meta


def parse_pdf(raw: bytes, filename: str) -> tuple[str, dict]:
    reader = PdfReader(io.BytesIO(raw))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        if txt.strip():
            pages.append(f"--- page {i + 1} ---\n{txt}")
    body = f"# {filename}\nPages: {len(reader.pages)}\n\n" + "\n\n".join(pages)
    return _truncate(body), {"page_count": len(reader.pages)}


def parse_docx(raw: bytes, filename: str) -> tuple[str, dict]:
    doc = _Docx(io.BytesIO(raw))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Tables — render as pipe-delimited lines.
    table_text: list[str] = []
    for ti, table in enumerate(doc.tables):
        table_text.append(f"--- table {ti + 1} ---")
        for row in table.rows:
            table_text.append(" | ".join(c.text.strip() for c in row.cells))
    body_parts = [f"# {filename}", f"Paragraphs: {len(paragraphs)}", *paragraphs]
    if table_text:
        body_parts.append("\n".join(table_text))
    return _truncate("\n".join(body_parts)), {
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
    }


def parse_txt(raw: bytes, filename: str) -> tuple[str, dict]:
    try:
        body = raw.decode("utf-8")
    except UnicodeDecodeError:
        body = raw.decode("utf-8", errors="replace")
    return _truncate(f"# {filename}\n\n{body}"), {"chars": len(body)}


def parse(raw: bytes, filename: str) -> tuple[str, str, dict]:
    """Return (kind, content_text, metadata). On parse error, falls back to txt."""
    kind = kind_for(filename)
    try:
        if kind == "csv":
            text, meta = parse_csv(raw, filename)
        elif kind == "xlsx":
            text, meta = parse_xlsx(raw, filename)
        elif kind == "pdf":
            text, meta = parse_pdf(raw, filename)
        elif kind == "docx":
            text, meta = parse_docx(raw, filename)
        else:
            text, meta = parse_txt(raw, filename)
    except Exception as e:
        # Don't fail the upload — store the raw bytes as txt and surface the
        # parse error in metadata so the user can see what happened.
        kind = "txt"
        text, meta = parse_txt(raw, filename)
        meta = {**meta, "parse_error": str(e)}
    return kind, text, meta
